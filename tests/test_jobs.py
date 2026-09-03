"""Тесты задач сравнения и пайплайна (spec: comparison-jobs)."""

import io
from types import SimpleNamespace

import pytest
from docx import Document

from app import create_app
from app.services import jobs, uploads


class MockChat:
    def __init__(self, responses, on_invoke=None):
        self.responses = list(responses)
        self.calls = 0
        self.on_invoke = on_invoke

    def invoke(self, messages):
        self.calls += 1
        if self.on_invoke:
            self.on_invoke()
        item = self.responses.pop(0)
        if isinstance(item, Exception):
            raise item
        return SimpleNamespace(content=item)


def docx_bytes(paragraphs):
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@pytest.fixture(autouse=True)
def clean_registries():
    uploads.clear_uploads()
    jobs.clear_jobs()


@pytest.fixture()
def make_app(tmp_path):
    def factory(chat):
        return create_app(
            {
                "TESTING": True,
                "UPLOAD_DIR": str(tmp_path / "uploads"),
                "JOBS_SYNCHRONOUS": True,
                "LLM_CHAT": chat,
            }
        )

    return factory


def docx_bytes_with_table(rows_data):
    doc = Document()
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    for i, row in enumerate(rows_data):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def upload_table_docx(client, name, rows_data):
    response = client.post(
        "/api/upload",
        data={"file": (docx_bytes_with_table(rows_data), name)},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    return response.get_json()["upload_id"]


def upload_docx(client, name, paragraphs):
    response = client.post(
        "/api/upload",
        data={"file": (docx_bytes(paragraphs), name)},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    return response.get_json()["upload_id"]


class TestStartCompare:
    def test_start_returns_202_and_job_id(self, make_app):
        app = make_app(MockChat(['{"label": "changed"}']))
        client = app.test_client()
        id1 = upload_docx(client, "v1.docx", ["А", "Б"])
        id2 = upload_docx(client, "v2.docx", ["А", "Х"])

        response = client.post(
            "/api/compare", json={"upload_id_1": id1, "upload_id_2": id2}
        )
        assert response.status_code == 202
        assert response.get_json()["job_id"]

    def test_unknown_upload_id_returns_404(self, make_app):
        client = make_app(MockChat([])).test_client()
        response = client.post(
            "/api/compare",
            json={"upload_id_1": "no-such", "upload_id_2": "no-such"},
        )
        assert response.status_code == 404
        assert "error" in response.get_json()

    def test_missing_ids_returns_400(self, make_app):
        client = make_app(MockChat([])).test_client()
        for payload in ({}, {"upload_id_1": "x"}, {"upload_id_2": "y"}):
            response = client.post("/api/compare", json=payload)
            assert response.status_code == 400
            assert "error" in response.get_json()


class TestJobStatus:
    def test_processing_status_with_stage_message(self, make_app):
        # Дано задача на этапе поиска различий
        client = make_app(MockChat([])).test_client()
        job_id = jobs.create_job()
        jobs.set_stage(job_id, "diffing")

        # Когда клиент опрашивает статус
        response = client.get(f"/api/jobs/{job_id}")
        # То статус «в обработке», ключ этапа и сообщение этапа на русском
        body = response.get_json()
        assert response.status_code == 200
        assert body["status"] == "processing"
        assert body["stage"] == "diffing"
        assert body["stage_message"] == "Поиск различий..."

    def test_unknown_job_returns_404(self, make_app):
        client = make_app(MockChat([])).test_client()
        response = client.get("/api/jobs/no-such-job")
        assert response.status_code == 404
        assert "error" in response.get_json()

    def test_done_job_returns_aligned_result(self, make_app):
        # Дано два документа с одним изменённым абзацем
        app = make_app(MockChat(['{"label": "changed"}']))
        client = app.test_client()
        id1 = upload_docx(client, "v1.docx", ["А", "ББ текст первый", "В"])
        id2 = upload_docx(client, "v2.docx", ["А", "ББ текст второй", "В"])
        job_id = client.post(
            "/api/compare", json={"upload_id_1": id1, "upload_id_2": id2}
        ).get_json()["job_id"]

        # Когда задача завершена
        body = client.get(f"/api/jobs/{job_id}").get_json()
        # То результат содержит выровненные строки с классами различий
        assert body["status"] == "done"
        result = body["result"]
        assert result["semantic"] is True
        rows = result["rows"]
        assert rows[0] == {
            "left": {"text": "А", "change": None},
            "right": {"text": "А", "change": None},
        }
        assert rows[1] == {
            "left": {
                "text": "ББ текст первый",
                "change": "changed",
                "segments": [
                    {"text": "ББ текст ", "type": "same"},
                    {"text": "первый", "type": "del"},
                ],
            },
            "right": {
                "text": "ББ текст второй",
                "change": "changed",
                "segments": [
                    {"text": "ББ текст ", "type": "same"},
                    {"text": "второй", "type": "add"},
                ],
            },
        }
        assert rows[2] == {
            "left": {"text": "В", "change": None},
            "right": {"text": "В", "change": None},
        }

    def test_removed_block_has_placeholder_on_right(self, make_app):
        app = make_app(MockChat(['{"label": "removed"}']))
        client = app.test_client()
        id1 = upload_docx(client, "v1.docx", ["А", "Б"])
        id2 = upload_docx(client, "v2.docx", ["А"])
        job_id = client.post(
            "/api/compare", json={"upload_id_1": id1, "upload_id_2": id2}
        ).get_json()["job_id"]

        rows = client.get(f"/api/jobs/{job_id}").get_json()["result"]["rows"]
        assert rows[1] == {
            "left": {"text": "Б", "change": "removed"},
            "right": None,
        }

    def test_added_block_has_placeholder_on_left(self, make_app):
        app = make_app(MockChat(['{"label": "added"}']))
        client = app.test_client()
        id1 = upload_docx(client, "v1.docx", ["А"])
        id2 = upload_docx(client, "v2.docx", ["А", "Б"])
        job_id = client.post(
            "/api/compare", json={"upload_id_1": id1, "upload_id_2": id2}
        ).get_json()["job_id"]

        rows = client.get(f"/api/jobs/{job_id}").get_json()["result"]["rows"]
        assert rows[1] == {
            "left": None,
            "right": {"text": "Б", "change": "added"},
        }

    def test_changed_removed_added_blocks_classified_separately(self, make_app):
        # Дано документ с изменённым, удалённым и добавленным абзацами подряд
        app = make_app(
            MockChat(
                ['{"label": "changed"}', '{"label": "removed"}', '{"label": "added"}']
            )
        )
        client = app.test_client()
        id1 = upload_docx(
            client,
            "v1.docx",
            [
                "Альфа, начало документа.",
                "Пункт второй: срок действия один год.",
                "Пункт третий: ответственность сторон по договору.",
                "Омега, конец документа.",
            ],
        )
        id2 = upload_docx(
            client,
            "v2.docx",
            [
                "Альфа, начало документа.",
                "Пункт второй: срок действия два года.",
                "Совершенно новый пункт про форс-мажор.",
                "Омега, конец документа.",
            ],
        )
        job_id = client.post(
            "/api/compare", json={"upload_id_1": id1, "upload_id_2": id2}
        ).get_json()["job_id"]

        # То изменения различаются: зелёный — изменён, красный — удалён, жёлтый — добавлен
        body = client.get(f"/api/jobs/{job_id}").get_json()
        assert body["status"] == "done"
        rows = body["result"]["rows"]
        assert rows[0]["left"]["change"] is None
        assert rows[1] == {
            "left": {
                "text": "Пункт второй: срок действия один год.",
                "change": "changed",
                "segments": [
                    {"text": "Пункт второй: срок действия ", "type": "same"},
                    {"text": "один", "type": "del"},
                    {"text": " ", "type": "same"},
                    {"text": "год", "type": "chg"},
                    {"text": ".", "type": "same"},
                ],
            },
            "right": {
                "text": "Пункт второй: срок действия два года.",
                "change": "changed",
                "segments": [
                    {"text": "Пункт второй: срок действия ", "type": "same"},
                    {"text": "два", "type": "add"},
                    {"text": " ", "type": "same"},
                    {"text": "года", "type": "chg"},
                    {"text": ".", "type": "same"},
                ],
            },
        }
        assert rows[2] == {
            "left": {
                "text": "Пункт третий: ответственность сторон по договору.",
                "change": "removed",
            },
            "right": None,
        }
        assert rows[3] == {
            "left": None,
            "right": {
                "text": "Совершенно новый пункт про форс-мажор.",
                "change": "added",
            },
        }
        assert rows[4]["left"]["change"] is None

    def test_table_rows_have_structured_cells(self, make_app):
        # Дано документы с таблицей: изменена одна ячейка
        app = make_app(MockChat(['{"label": "changed"}']))
        client = app.test_client()
        id1 = upload_table_docx(client, "v1.docx", [["Товар", "Цена"], ["Яблоки", "100"]])
        id2 = upload_table_docx(client, "v2.docx", [["Товар", "Цена"], ["Яблоки", "150"]])
        job_id = client.post(
            "/api/compare", json={"upload_id_1": id1, "upload_id_2": id2}
        ).get_json()["job_id"]

        rows = client.get(f"/api/jobs/{job_id}").get_json()["result"]["rows"]
        # Шапка: структурные ячейки, без изменений
        assert rows[0]["left"]["cells"] == ["Товар", "Цена"]
        assert rows[0]["left"]["change"] is None
        # Служебная строка-разделитель помечена sep
        assert rows[1]["left"]["sep"] is True
        assert "cells" not in rows[1]["left"]
        # Изменённая строка: ячейки + пословный diff по ячейкам
        assert rows[2]["left"]["cells"] == ["Яблоки", "100"]
        assert rows[2]["left"]["cell_segments"] == [
            [{"text": "Яблоки", "type": "same"}],
            [{"text": "100", "type": "chg"}],
        ]
        assert rows[2]["right"]["cell_segments"] == [
            [{"text": "Яблоки", "type": "same"}],
            [{"text": "150", "type": "chg"}],
        ]

    def test_table_with_different_column_counts(self, make_app):
        # Дано в файле 2 у таблицы добавилась колонка
        app = make_app(
            MockChat(['{"label": "changed"}', '{"label": "changed"}'])
        )
        client = app.test_client()
        id1 = upload_table_docx(client, "v1.docx", [["A", "B"], ["1", "2"]])
        id2 = upload_table_docx(client, "v2.docx", [["A", "B", "C"], ["1", "2", "3"]])
        job_id = client.post(
            "/api/compare", json={"upload_id_1": id1, "upload_id_2": id2}
        ).get_json()["job_id"]

        body = client.get(f"/api/jobs/{job_id}").get_json()
        assert body["status"] == "done"
        rows = body["result"]["rows"]
        # Строка данных: правая сторона имеет 3 ячейки, левая — 2,
        # пословный diff по ячейкам выровнен (новая колонка — «добавлено»)
        data_row = next(r for r in rows if r["left"] and r["left"].get("cells") == ["1", "2"])
        assert data_row["right"]["cells"] == ["1", "2", "3"]
        assert len(data_row["left"]["cell_segments"]) == 3
        assert len(data_row["right"]["cell_segments"]) == 3
        assert data_row["right"]["cell_segments"][2] == [{"text": "3", "type": "add"}]
        assert data_row["left"]["cell_segments"][2] == []

    def test_failed_job_returns_error(self, make_app):
        # Дано файл с расширением .docx, но битым содержимым
        app = make_app(MockChat([]))
        client = app.test_client()
        bad = client.post(
            "/api/upload",
            data={"file": (io.BytesIO(b"not a docx"), "bad.docx")},
            content_type="multipart/form-data",
        ).get_json()["upload_id"]
        good = upload_docx(client, "good.docx", ["А"])

        job_id = client.post(
            "/api/compare", json={"upload_id_1": bad, "upload_id_2": good}
        ).get_json()["job_id"]

        # То задача завершается ошибкой с понятным сообщением
        body = client.get(f"/api/jobs/{job_id}").get_json()
        assert body["status"] == "failed"
        assert body["error"]

    def test_llm_unavailable_degrades_to_opcode_classification(self, make_app):
        # Дано LLM недоступна
        app = make_app(MockChat([ConnectionError("down")]))
        client = app.test_client()
        id1 = upload_docx(client, "v1.docx", ["А", "ББ текст первый"])
        id2 = upload_docx(client, "v2.docx", ["А", "ББ текст второй"])
        job_id = client.post(
            "/api/compare", json={"upload_id_1": id1, "upload_id_2": id2}
        ).get_json()["job_id"]

        # То задача всё равно завершается успешно с пометкой деградации
        body = client.get(f"/api/jobs/{job_id}").get_json()
        assert body["status"] == "done"
        assert body["result"]["semantic"] is False
        assert body["result"]["rows"][1]["left"]["change"] == "changed"

    def test_llm_stage_during_classification(self, make_app):
        # Дано задача запущена; во время вызова LLM этап должен быть «Анализ через LLM...»
        observed = {}

        def on_invoke():
            job = next(iter(jobs._JOBS.values()))
            observed["stage"] = job["stage_message"]

        app = make_app(MockChat(['{"label": "changed"}'], on_invoke=on_invoke))
        client = app.test_client()
        id1 = upload_docx(client, "v1.docx", ["А"])
        id2 = upload_docx(client, "v2.docx", ["Б"])
        client.post("/api/compare", json={"upload_id_1": id1, "upload_id_2": id2})

        assert observed["stage"] == "Анализ через LLM..."
