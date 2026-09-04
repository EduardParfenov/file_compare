"""Тесты конвертации .docx в Markdown (spec: markdown-conversion)."""

import pytest
from docx import Document

from app.services.docx_converter import ConversionError, convert_docx


def make_docx(path, builder):
    doc = Document()
    builder(doc)
    doc.save(path)
    return str(path)


def test_headings_and_paragraphs(tmp_path):
    # Дано .docx с заголовком 1-го уровня и абзацем
    path = make_docx(
        tmp_path / "doc.docx",
        lambda doc: (
            doc.add_heading("Отчёт", level=1),
            doc.add_paragraph("Текст отчёта"),
        ),
    )
    # Когда выполняется конвертация
    markdown = convert_docx(path)
    # То заголовок становится Markdown-заголовком, абзац — текстовым блоком
    assert markdown == "# Отчёт\n\nТекст отчёта"


def test_heading_levels(tmp_path):
    path = make_docx(
        tmp_path / "doc.docx",
        lambda doc: (
            doc.add_heading("Раздел", level=2),
            doc.add_paragraph("Текст"),
        ),
    )
    markdown = convert_docx(path)
    assert markdown.startswith("## Раздел")


def test_table(tmp_path):
    # Дано .docx с таблицей 2x2
    def build(doc):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"

    path = make_docx(tmp_path / "doc.docx", build)
    # Когда выполняется конвертация
    markdown = convert_docx(path)
    # То результат содержит Markdown-таблицу с теми же ячейками
    lines = markdown.splitlines()
    assert "| A | B |" in lines
    assert "| 1 | 2 |" in lines
    assert any(set(line) <= set("|- ") for line in lines), "нет строки-разделителя"


def test_merged_cells_flattened(tmp_path):
    # Дано таблица 3x3: горизонтальный merge в строке 0 (колонки 0-1),
    # вертикальный merge в колонке 2 (строки 1-2)
    def build(doc):
        table = doc.add_table(rows=3, cols=3)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"R{i}C{j}"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(1, 2).merge(table.cell(2, 2))

    path = make_docx(tmp_path / "doc.docx", build)
    # Когда выполняется конвертация
    lines = convert_docx(path).splitlines()
    # То текст объединённой ячейки эмитируется один раз, повторы — пустые,
    # а переносы внутри ячейки сворачиваются в пробел
    assert lines[0] == "| R0C0 R0C1 |  | R0C2 |"
    assert lines[2] == "| R1C0 | R1C1 | R1C2 R2C2 |"
    assert lines[3] == "| R2C0 | R2C1 |  |"
    # И все строки таблицы — однострочные (без разрывов от \n в ячейках)
    assert all("\n" not in line for line in lines)


def test_multiline_cell_becomes_one_line(tmp_path):
    # Дано ячейка таблицы с переносом строки внутри
    def build(doc):
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "строка 1\nстрока 2"
        table.cell(0, 1).text = "B"

    path = make_docx(tmp_path / "doc.docx", build)
    # Когда выполняется конвертация
    markdown = convert_docx(path)
    # То ячейка не разрывает Markdown-строку таблицы
    assert "| строка 1 строка 2 | B |" in markdown.splitlines()


def test_empty_document(tmp_path):
    path = make_docx(tmp_path / "doc.docx", lambda doc: None)
    assert convert_docx(path) == ""


def test_corrupted_file(tmp_path):
    # Дано файл с расширением .docx, не являющийся валидным DOCX
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a real docx")
    # Когда/То конвертация завершается понятной ошибкой
    with pytest.raises(ConversionError):
        convert_docx(str(path))


def test_deterministic(tmp_path):
    path = make_docx(
        tmp_path / "doc.docx", lambda doc: doc.add_paragraph("Текст")
    )
    assert convert_docx(path) == convert_docx(path)
